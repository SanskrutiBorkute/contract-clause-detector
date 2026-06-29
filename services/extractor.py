# ClauseGuard AI — Text Extraction Service
import os
import pdfplumber

class TextExtractor:
    @staticmethod
    def extract_from_txt(file_path):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    @staticmethod
    def extract_from_docx(file_path):
        try:
            import docx
        except ImportError:
            raise ImportError(
                "python-docx package is required for Word document analysis. "
                "Please run: pip install python-docx"
            )
        
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return "\n".join(full_text)

    @staticmethod
    def extract_from_pdf(file_path):
        text = ""
        # 1. Try pypdfium2 first for ultra-fast C++ based text extraction (critical for 300-page performance)
        try:
            import pypdfium2 as pdfium
            doc = pdfium.PdfDocument(file_path)
            for page in doc:
                textpage = page.get_textpage()
                page_text = textpage.get_text_range()
                if page_text:
                    text += page_text + "\n"
            print("PDF extracted successfully via pypdfium2 (fast).")
        except Exception as e:
            print(f"pypdfium2 fast extraction failed: {e}. Trying pdfplumber...")
            text = ""

        # 2. Fallback to pdfplumber if pypdfium2 failed or returned empty
        if not text.strip():
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                print("PDF extracted successfully via pdfplumber.")
            except Exception as e:
                print(f"pdfplumber extraction failed: {e}")
        
        # 3. If still empty, treat as scanned PDF and perform OCR
        if not text.strip():
            print("PDF extracted text was empty. Attempting OCR fallback...")
            return TextExtractor.extract_from_scanned_pdf(file_path)
        
        return text

    @staticmethod
    def extract_from_scanned_pdf(file_path):
        try:
            import pypdfium2 as pdfium
            from PIL import Image
        except ImportError:
            return ""

        text = ""
        try:
            doc = pdfium.PdfDocument(file_path)
            # Limit OCR to first 10 pages for performance if it's a huge scanned document
            max_pages = min(len(doc), 10)
            print(f"Performing OCR on first {max_pages} pages of scanned PDF...")
            for i in range(max_pages):
                page = doc[i]
                bitmap = page.render(scale=2) # render to high-res image
                pil_img = bitmap.to_pil()
                page_text = TextExtractor.perform_ocr_on_image(pil_img)
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            print(f"Scanned PDF OCR failed: {e}")
        
        return text

    @staticmethod
    def extract_from_image(file_path):
        from PIL import Image
        try:
            img = Image.open(file_path)
            return TextExtractor.perform_ocr_on_image(img)
        except Exception as e:
            raise Exception(f"Image opening/processing failed: {e}")

    @staticmethod
    def perform_ocr_on_image(pil_image):
        # Try EasyOCR first
        try:
            import numpy as np
            import easyocr
            img_np = np.array(pil_image)
            reader = easyocr.Reader(['en'], gpu=False)
            results = reader.readtext(img_np)
            return "\n".join([r[1] for r in results])
        except Exception as e:
            print(f"EasyOCR failed: {e}. Trying pytesseract fallback...")
        
        # Fallback to pytesseract
        try:
            import pytesseract
            return pytesseract.image_to_string(pil_image)
        except Exception as e:
            print(f"Pytesseract failed: {e}")
        
        raise Exception(
            "OCR engines (EasyOCR and Pytesseract) are unavailable. "
            "Please ensure Tesseract OCR is installed or numpy/torch versions are compatible."
        )

    @classmethod
    def get_page_count(cls, file_path):
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            return max(1, len(text) // 2500)
        elif ext == ".docx":
            try:
                import docx
                doc = docx.Document(file_path)
                text_len = sum(len(p.text) for p in doc.paragraphs)
                return max(1, text_len // 2500)
            except Exception:
                return 1
        elif ext == ".pdf":
            # Try pypdfium2 first
            try:
                import pypdfium2 as pdfium
                doc = pdfium.PdfDocument(file_path)
                return len(doc)
            except Exception:
                # Fallback to pdfplumber
                try:
                    with pdfplumber.open(file_path) as pdf:
                        return len(pdf.pages)
                except Exception:
                    return 1
        elif ext in [".png", ".jpg", ".jpeg"]:
            return 1
        return 1

    @classmethod
    def extract_text(cls, file_path):
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".txt":
            return cls.extract_from_txt(file_path)
        elif ext == ".docx":
            return cls.extract_from_docx(file_path)
        elif ext == ".pdf":
            return cls.extract_from_pdf(file_path)
        elif ext in [".png", ".jpg", ".jpeg"]:
            return cls.extract_from_image(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

